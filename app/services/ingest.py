from __future__ import annotations

import csv
import hashlib
import io
import json
import re
import secrets
from collections import Counter
from datetime import UTC, datetime
from pathlib import Path
from statistics import mean

import numpy as np
from dateutil import parser as date_parser
from fastapi import UploadFile
from pypdf import PdfReader
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.core.metrics import INGESTED_DOCUMENTS
from app.models import Document, DocumentStatus
from app.services.storage import ArtifactStore, write_json

try:
    from docx import Document as DocxDocument
except Exception:  # pragma: no cover
    DocxDocument = None


EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
PHONE_RE = re.compile(r"(?:\+?\d{1,2}[\s.-]?)?(?:\(?\d{3}\)?[\s.-]?)\d{3}[\s.-]?\d{4}")
SSN_RE = re.compile(r"\b\d{3}-\d{2}-\d{4}\b")
CARD_RE = re.compile(r"\b(?:\d[ -]*?){13,16}\b")


class IngestionService:
    def __init__(self, db: Session):
        self.db = db
        self.settings = get_settings()
        self.store = ArtifactStore()

    def ingest_upload(
        self,
        *,
        tenant_id: str,
        project_id: str,
        file: UploadFile,
        metadata: dict,
    ) -> Document:
        if not isinstance(metadata, dict):
            raise ValueError("Metadata must be an object")
        if len(json.dumps(metadata)) > 32_000:
            raise ValueError("Metadata payload too large")

        filename = self._safe_filename(file.filename or "uploaded")
        content = file.file.read()
        if not content:
            raise ValueError("Uploaded file is empty")

        file_type = self._detect_file_type(filename)
        self._mock_virus_scan(filename, content)

        extracted_text, extraction_meta = self._extract_text(content, file_type)
        normalized_text = self._normalize_text(extracted_text)
        sections = self._extract_sections(normalized_text)
        sha_hash = hashlib.sha256(normalized_text.encode("utf-8")).hexdigest()

        exact_duplicate = self.db.scalar(
            select(Document).where(
                Document.tenant_id == tenant_id,
                Document.project_id == project_id,
                Document.sha256_hash == sha_hash,
            )
        )

        raw_path = self.store.raw_docs_dir(tenant_id, project_id) / f"{secrets.token_hex(8)}_{filename}"
        raw_path.write_bytes(content)

        normalized_path = self.store.normalized_dir(tenant_id, project_id) / f"{sha_hash}.json"
        normalized_payload = {
            "filename": filename,
            "text": normalized_text,
            "sections": sections,
            "metadata": metadata,
            "extraction": extraction_meta,
        }
        write_json(normalized_path, normalized_payload)

        pii_hits = self._detect_pii(normalized_text)
        near_duplicate_of = None
        similarity_score = 0.0
        if not exact_duplicate:
            near_duplicate_of, similarity_score = self._find_near_duplicate(
                tenant_id=tenant_id,
                project_id=project_id,
                normalized_text=normalized_text,
            )

        quality_score = self._doc_quality_score(
            text=normalized_text,
            extraction_meta=extraction_meta,
            metadata=metadata,
            pii_hits=pii_hits,
            similarity_score=similarity_score,
        )

        status = DocumentStatus.READY
        if exact_duplicate:
            status = DocumentStatus.REJECTED
            near_duplicate_of = exact_duplicate.id
        elif pii_hits:
            status = DocumentStatus.REDACTION_REQUIRED
        elif quality_score < self.settings.doc_quality_threshold or near_duplicate_of:
            status = DocumentStatus.NEEDS_REVIEW

        doc = Document(
            tenant_id=tenant_id,
            project_id=project_id,
            filename=filename,
            file_type=file_type,
            storage_path=str(raw_path),
            normalized_text_path=str(normalized_path),
            sha256_hash=sha_hash,
            near_duplicate_of=near_duplicate_of,
            quality_score=quality_score,
            pii_hits=pii_hits,
            metadata_json=metadata,
            status=status,
        )

        self.db.add(doc)
        self.db.commit()
        self.db.refresh(doc)
        INGESTED_DOCUMENTS.labels(status=doc.status.value).inc()
        return doc

    @staticmethod
    def _mock_virus_scan(filename: str, content: bytes) -> None:
        settings = get_settings()
        if len(content) > settings.max_upload_mb * 1024 * 1024:
            raise ValueError(f"File too large: {filename}")

    @staticmethod
    def _safe_filename(filename: str) -> str:
        base = Path(filename).name
        sanitized = re.sub(r"[^a-zA-Z0-9._-]", "_", base)
        sanitized = sanitized.strip("._")
        return sanitized or "uploaded"

    @staticmethod
    def _detect_file_type(filename: str) -> str:
        suffix = Path(filename).suffix.lower()
        mapping = {
            ".pdf": "pdf",
            ".docx": "docx",
            ".txt": "txt",
            ".md": "markdown",
            ".html": "html",
            ".htm": "html",
            ".csv": "csv",
        }
        if suffix not in mapping:
            raise ValueError(f"Unsupported file type: {suffix}")
        return mapping[suffix]

    def _extract_text(self, content: bytes, file_type: str) -> tuple[str, dict]:
        if file_type in {"txt", "markdown"}:
            return content.decode("utf-8", errors="ignore"), {"ocr_used": False, "ocr_confidence": 1.0}
        if file_type == "html":
            raw = content.decode("utf-8", errors="ignore")
            cleaned = re.sub(r"<[^>]+>", " ", raw)
            return cleaned, {"ocr_used": False, "ocr_confidence": 1.0}
        if file_type == "csv":
            raw = content.decode("utf-8", errors="ignore")
            reader = csv.reader(io.StringIO(raw))
            lines = [" | ".join(row) for row in reader]
            return "\n".join(lines), {"ocr_used": False, "ocr_confidence": 1.0}
        if file_type == "pdf":
            reader = PdfReader(io.BytesIO(content))
            pages = []
            for page in reader.pages:
                pages.append(page.extract_text() or "")
            text = "\n".join(pages).strip()
            if not text:
                return "", {"ocr_used": True, "ocr_confidence": 0.2}
            return text, {"ocr_used": False, "ocr_confidence": 0.95}
        if file_type == "docx":
            if DocxDocument is None:
                raise ValueError("python-docx not installed")
            document = DocxDocument(io.BytesIO(content))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            return text, {"ocr_used": False, "ocr_confidence": 1.0}
        raise ValueError(f"Unhandled file type: {file_type}")

    @staticmethod
    def _normalize_text(text: str) -> str:
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"[\t ]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    @staticmethod
    def _extract_sections(text: str) -> list[dict]:
        if not text:
            return []
        sections: list[dict] = []
        current_title = "General"
        current_lines: list[str] = []

        for line in text.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.endswith(":") or stripped.startswith("#") or re.match(r"^\d+\.\s", stripped):
                if current_lines:
                    sections.append({"title": current_title, "content": " ".join(current_lines)})
                    current_lines = []
                current_title = stripped.strip("# ")
            else:
                current_lines.append(stripped)

        if current_lines:
            sections.append({"title": current_title, "content": " ".join(current_lines)})
        return sections

    @staticmethod
    def _detect_pii(text: str) -> list[dict]:
        hits: list[dict] = []
        for label, regex in (
            ("email", EMAIL_RE),
            ("phone", PHONE_RE),
            ("ssn", SSN_RE),
            ("card", CARD_RE),
        ):
            for match in regex.findall(text):
                hits.append({"type": label, "value": str(match)[:24]})
        return hits[:100]

    def _find_near_duplicate(
        self,
        *,
        tenant_id: str,
        project_id: str,
        normalized_text: str,
    ) -> tuple[str | None, float]:
        target_vec = self._hashed_embedding(normalized_text)
        candidate_docs = list(
            self.db.scalars(
                select(Document).where(
                    Document.tenant_id == tenant_id,
                    Document.project_id == project_id,
                    Document.status != DocumentStatus.REJECTED,
                )
            ).all()
        )

        best_doc_id = None
        best_score = 0.0
        for doc in candidate_docs:
            payload = json.loads(Path(doc.normalized_text_path).read_text(encoding="utf-8"))
            candidate_text = self._normalize_text(payload.get("text", ""))
            if not candidate_text:
                continue
            score = self._cosine_similarity(target_vec, self._hashed_embedding(candidate_text))
            if score > best_score:
                best_score = score
                best_doc_id = doc.id

        if best_score >= self.settings.near_duplicate_threshold:
            return best_doc_id, best_score
        return None, best_score

    @staticmethod
    def _tokenize(text: str) -> list[str]:
        return re.findall(r"[a-zA-Z0-9]{2,}", text.lower())

    def _hashed_embedding(self, text: str, dim: int = 256) -> np.ndarray:
        vector = np.zeros(dim, dtype=np.float32)
        tokens = self._tokenize(text)
        if not tokens:
            return vector
        counts = Counter(tokens)
        for token, count in counts.items():
            digest = hashlib.sha1(token.encode("utf-8")).hexdigest()
            index = int(digest[:8], 16) % dim
            vector[index] += float(count)
        norm = np.linalg.norm(vector)
        if norm == 0:
            return vector
        return vector / norm

    @staticmethod
    def _cosine_similarity(a: np.ndarray, b: np.ndarray) -> float:
        if a.size == 0 or b.size == 0:
            return 0.0
        denom = float(np.linalg.norm(a) * np.linalg.norm(b))
        if denom == 0:
            return 0.0
        return float(np.dot(a, b) / denom)

    def _doc_quality_score(
        self,
        *,
        text: str,
        extraction_meta: dict,
        metadata: dict,
        pii_hits: list[dict],
        similarity_score: float,
    ) -> int:
        if not text:
            return 0

        extraction_quality = 100
        if extraction_meta.get("ocr_used"):
            extraction_quality = int(extraction_meta.get("ocr_confidence", 0.0) * 100)
        printable_ratio = self._printable_ratio(text)
        extraction_quality = int((extraction_quality + printable_ratio) / 2)

        structure_quality = min(100, int(30 + self._heading_count(text) * 10 + self._bullet_count(text) * 4))

        freshness = self._freshness_score(metadata)

        redaction_risk = max(0, 100 - len(pii_hits) * 20)

        dedupe_penalty = int((1.0 - similarity_score) * 100)

        components = [extraction_quality, structure_quality, freshness, redaction_risk, dedupe_penalty]
        return int(mean(components))

    @staticmethod
    def _printable_ratio(text: str) -> int:
        if not text:
            return 0
        printable = sum(1 for c in text if c.isprintable())
        return int((printable / len(text)) * 100)

    @staticmethod
    def _heading_count(text: str) -> int:
        return sum(
            1
            for line in text.splitlines()
            if line.strip().startswith(("#", "1.", "2.", "3.")) or line.strip().endswith(":")
        )

    @staticmethod
    def _bullet_count(text: str) -> int:
        return sum(1 for line in text.splitlines() if line.strip().startswith(("-", "*")))

    @staticmethod
    def _freshness_score(metadata: dict) -> int:
        effective_date = metadata.get("effective_date")
        if not effective_date:
            return 60
        try:
            dt = date_parser.parse(str(effective_date))
        except Exception:
            return 40
        days_old = (datetime.now(tz=UTC) - dt.astimezone(UTC)).days
        if days_old <= 180:
            return 100
        if days_old <= 365:
            return 85
        if days_old <= 730:
            return 70
        return 45
